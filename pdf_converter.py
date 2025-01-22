# 基础库
import os
import io
import numpy as np
import cv2
from PIL import Image

# GUI相关
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from threading import Thread
import queue

# PDF处理
import PyPDF2
import fitz  # PyMuPDF库

# 文档处理
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# OCR相关
import pytesseract
import easyocr
import torch
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import onnxruntime as ort
from onnxruntime.capi.onnxruntime_pybind11_state import InvalidGraph
from offline_model import setup_offline_model

class PDFConverter:
    def __init__(self):
        self.pdf_file = None
        self.output_file = None
        self.progress_callback = None
        
        # 检测可用的计算设备
        if torch.cuda.is_available():
            self.device = 'cuda'
        elif hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
            self.device = 'mps'  # 对于 Apple Silicon
        else:
            self.device = 'cpu'
        print(f"PyTorch使用设备: {self.device}")
        
        # 初始化离线OCR模型
        try:
            self.processor, self.model = setup_offline_model()
            if self.processor and self.model:
                self.model = self.model.to(self.device)
                self.trocr_available = True
                print("离线模型加载成功")
            else:
                raise Exception("离线模型加载失败")
        except Exception as e:
            print(f"模型初始化失败: {str(e)}")
            self.trocr_available = False

        # 检测 ONNX Runtime 可用的提供程序
        try:
            providers = ort.get_available_providers()
            print(f"可用的计算提供程序: {providers}")
            
            self.available_providers = []
            # 按优先级排序提供程序
            if 'NPUExecutionProvider' in providers:
                self.available_providers.append('NPUExecutionProvider')
                print("使用 NPU 加速")
            elif 'DmlExecutionProvider' in providers:
                self.available_providers.append('DmlExecutionProvider')
                print("使用 DirectML 加速")
            elif 'CUDAExecutionProvider' in providers:
                self.available_providers.append('CUDAExecutionProvider')
                print("使用 CUDA 加速")
            
            # 始终添加 CPU 作为后备选项
            self.available_providers.append('CPUExecutionProvider')
            
        except Exception as e:
            print(f"初始化计算设备失败: {str(e)}")
            self.available_providers = ['CPUExecutionProvider']

        # 初始化 EasyOCR
        try:
            # 根据可用的加速方案设置 gpu 参数
            gpu_enabled = any(provider in providers for provider in 
                            ['NPUExecutionProvider', 'DmlExecutionProvider', 'CUDAExecutionProvider'])
            self.reader = easyocr.Reader(['ch_sim', 'en'], gpu=gpu_enabled)  # 只加载简体中文和英语
            self.easyocr_available = True
            print(f"EasyOCR初始化成功，{'启用' if gpu_enabled else '未启用'}硬件加速")
        except Exception as e:
            print(f"EasyOCR初始化失败: {str(e)}")
            self.easyocr_available = False

    def set_progress_callback(self, callback):
        """设置进度回调函数"""
        self.progress_callback = callback
    
    def update_progress(self, value):
        """更新进度"""
        if self.progress_callback:
            self.progress_callback(value)

    def ocr_with_trocr(self, image):
        """使用TrOCR进行高精度识别"""
        try:
            if not isinstance(image, Image.Image):
                image = Image.fromarray(image)
            
            # 图像预处理
            processed_image = self.preprocess_image(image)

            # 预处理图像
            pixel_values = self.processor(processed_image, return_tensors="pt").pixel_values
            pixel_values = pixel_values.to(self.device)

            with torch.no_grad():  # 禁用梯度计算以节省内存
                # 使用模型进行预测
                generated_ids = self.model.generate(pixel_values)
                generated_text = self.processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
            
            return generated_text
        except Exception as e:
            print(f"TrOCR错误: {str(e)}")
            return ""

    def preprocess_image(self, image):
        """图像预处理以提高OCR效果"""
        try:
            # 转换为numpy数组
            if isinstance(image, Image.Image):
                image = np.array(image)
            
            # 转换为灰度图
            gray = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
            
            # 自适应阈值化
            binary = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # 降噪
            denoised = cv2.fastNlMeansDenoising(binary)
            
            # 提高对比度
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            
            # 进行形态学操作以增强数学符号
            kernel = np.ones((3, 3), np.uint8)
            morph = cv2.morphologyEx(enhanced, cv2.MORPH_CLOSE, kernel)

            return Image.fromarray(morph)
        except Exception as e:
            print(f"图像预处理错误: {str(e)}")
            return image

    def pdf_to_text_with_format(self, pdf_path):
        """将PDF转换为带格式的文本，保留排版和图片"""
        try:
            doc = fitz.open(pdf_path)
            formatted_text = []
            total_pages = len(doc)
            
            print(f"总页数: {total_pages}")  # 打印总页数
            
            for page_num, page in enumerate(doc):
                # 1. 提取文本块和图片
                blocks = page.get_text("dict")["blocks"]
                page_items = []
                
                for block in blocks:
                    # 处理图片块
                    if block.get("type") == 1:  # 图片类型
                        try:
                            xref = block["xref"]  # 图片引用号
                            image = doc.extract_image(xref)
                            if image:
                                # 直接转存图片
                                image_data = image["image"]
                                page_items.append({
                                    'type': 'image',
                                    'data': image_data,
                                    'bbox': block["bbox"],  # 图片位置信息
                                    'width': image["width"],
                                    'height': image["height"]
                                })
                                print(f"提取图片: {image['width']}x{image['height']}")  # 打印图片尺寸
                        except Exception as e:
                            print(f"图片提取错误: {str(e)}")
                    
                    # 处理文本块
                    elif "lines" in block:
                        current_paragraph = {
                            'type': 'text',
                            'spans': [],
                            'bbox': block["bbox"],
                            'alignment': self._get_alignment(block["bbox"], page.rect.width)
                        }
                        
                        for line in block["lines"]:
                            for span in line["spans"]:
                                current_paragraph['spans'].append({
                                    'text': span['text'],
                                    'font': span['font'],
                                    'size': span['size'],
                                    'color': span['color'],
                                    'flags': span.get('flags', 0)  # 字体样式（粗体、斜体等）
                                })
                        
                        page_items.append(current_paragraph)
                
                formatted_text.extend(page_items)
                self.update_progress((page_num + 1) / total_pages * 50)
            
            doc.close()
            return formatted_text
            
        except Exception as e:
            raise Exception(f"PDF处理错误: {str(e)}")
    
    def _get_alignment(self, bbox, page_width):
        """根据文本块位置判断对齐方式"""
        x0, _, x1, _ = bbox
        center = page_width / 2
        margin = 20  # 边距容差
        
        if abs(x0 - margin) < 10:  # 左对齐
            return 0
        elif abs(x1 - (page_width - margin)) < 10:  # 右对齐
            return 2
        elif abs((x0 + x1) / 2 - center) < 10:  # 居中对齐
            return 1
        return 0  # 默认左对齐
    
    def _analyze_paragraph_structure(self, blocks):
        """分析段落结构，合并相关文本块"""
        paragraphs = []
        current_paragraph = None
        
        for block in blocks:
            if "lines" not in block:
                continue
            
            # 获取块的位置和文本信息
            block_bbox = block["bbox"]
            block_text = ""
            block_style = {
                'font': None,
                'size': None,
                'color': None,
                'flags': 0
            }
            
            # 提取文本和样式
            for line in block["lines"]:
                for span in line["spans"]:
                    block_text += span["text"]
                    # 使用第一个span的样式作为主要样式
                    if not block_style['font']:
                        block_style = {
                            'font': span['font'],
                            'size': span['size'],
                            'color': span['color'],
                            'flags': span.get('flags', 0)
                        }
            
            # 判断是否是新段落的开始
            is_new_paragraph = True
            if current_paragraph:
                # 检查垂直距离
                prev_bbox = current_paragraph['bbox']
                vertical_gap = block_bbox[1] - prev_bbox[3]
                
                # 检查水平缩进
                indent = block_bbox[0] - prev_bbox[0]
                
                # 检查字体样式
                same_style = (
                    current_paragraph['style']['font'] == block_style['font'] and
                    abs(current_paragraph['style']['size'] - block_style['size']) < 0.5
                )
                
                # 如果垂直间距小且样式相同，认为是同一段落
                if vertical_gap < block_style['size'] * 1.5 and same_style:
                    is_new_paragraph = False
            
            if is_new_paragraph:
                if current_paragraph:
                    paragraphs.append(current_paragraph)
                
                # 创建新段落
                current_paragraph = {
                    'type': 'text',
                    'text': block_text,
                    'bbox': block_bbox,
                    'style': block_style,
                    'alignment': self._get_alignment(block_bbox, block["page_width"]),
                    'is_title': block_style['size'] > 12 or bool(block_style['flags'] & 2),  # 大字号或粗体可能是标题
                    'spans': []
                }
            else:
                # 合并到当前段落
                current_paragraph['text'] += '\n' + block_text
                current_paragraph['bbox'] = (
                    min(current_paragraph['bbox'][0], block_bbox[0]),
                    current_paragraph['bbox'][1],
                    max(current_paragraph['bbox'][2], block_bbox[2]),
                    block_bbox[3]
                )
            
            # 保存原始spans以保留详细格式
            for line in block["lines"]:
                current_paragraph['spans'].extend(line["spans"])
        
        if current_paragraph:
            paragraphs.append(current_paragraph)
        
        return paragraphs

    def _format_paragraph_in_docx(self, doc, paragraph_info):
        """在Word文档中格式化段落"""
        try:
            p = doc.add_paragraph()
            
            # 设置段落对齐方式
            p.alignment = paragraph_info['alignment']
            
            # 设置段落间距
            p.paragraph_format.space_before = Pt(12 if paragraph_info.get('is_title', False) else 6)
            p.paragraph_format.space_after = Pt(12 if paragraph_info.get('is_title', False) else 6)
            p.paragraph_format.line_spacing = 1.15
            
            # 处理段落缩进
            if not paragraph_info.get('is_title', False):
                p.paragraph_format.first_line_indent = Pt(21)  # 首行缩进2个字符
            
            # 添加文本并保持原始格式
            for span in paragraph_info['spans']:
                run = p.add_run(span['text'])
                
                # 统一设置字体为微软雅黑
                run.font.name = '微软雅黑'
                
                # 设置字体大小
                if span['size']:
                    try:
                        size = float(span['size'])
                        run.font.size = Pt(size)  # 保留原始大小
                    except:
                        run.font.size = Pt(11)  # 默认大小
                
                # 设置颜色
                if span['color']:
                    try:
                        if isinstance(span['color'], int):
                            color = span['color']
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                        else:
                            r, g, b = span['color'][:3]
                        run.font.color.rgb = RGBColor(r, g, b)
                    except:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                
                # 设置字体样式
                flags = span.get('flags', 0)
                run.font.bold = bool(flags & 2) or paragraph_info.get('is_title', False)
                run.font.italic = bool(flags & 1)
                run.font.underline = bool(flags & 4)
        
        except Exception as e:
            print(f"段落格式化错误: {str(e)}")

    def text_to_docx_with_format(self, formatted_text, output_path):
        """将带格式的文本转换为DOCX文档，保留排版和图片"""
        try:
            doc = Document()
            
            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = '微软雅黑'  # 统一字体为微软雅黑
            style.font.size = Pt(11)
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # 处理每个文本块和图片
            for item in formatted_text:
                if item['type'] == 'image':
                    # 处理图片
                    try:
                        image_stream = io.BytesIO(item['data'])
                        doc.add_picture(image_stream, 
                                        width=Inches(item['width'] / 72.0),
                                        height=Inches(item['height'] / 72.0))
                    except Exception as e:
                        print(f"图片插入错误: {str(e)}")
                
                elif item['type'] == 'text':
                    # 直接添加文本，不进行简体转换
                    self._format_paragraph_in_docx(doc, item)
            
            # 保存文档
            doc.save(output_path)
            print(f"成功创建文档: {output_path}")

            # 读取文档内容并进行简体转换
            for paragraph in doc.paragraphs:
                paragraph.text = self._convert_to_simplified_chinese(paragraph.text)

            # 重新保存文档
            doc.save(output_path)
            print(f"成功将文档转换为简体中文: {output_path}")
            
            return True
            
        except Exception as e:
            print(f"DOCX创建错误: {str(e)}")
            return False
    
    def convert_pdf_to_docx(self, input_pdf, output_docx):
        """将PDF转换为DOCX"""
        try:
            doc = Document()
            doc.add_heading('PDF 转换结果', level=1)

            # 读取PDF内容并添加到文档
            formatted_text = self.pdf_to_text_with_format(input_pdf)
            total_items = len(formatted_text)  # 获取总的内容项
            
            for i, item in enumerate(formatted_text):
                if item['type'] == 'text':
                    self._format_paragraph_in_docx(doc, item)
                elif item['type'] == 'image':
                    # 处理图片
                    image_stream = io.BytesIO(item['data'])
                    doc.add_picture(image_stream, 
                                    width=Inches(item['width']/72.0),
                                    height=Inches(item['height']/72.0))
                
                # 更新进度
                self.update_progress((i + 1) / total_items * 100)

            doc.save(output_docx)
            print(f"成功创建文档: {output_docx}")
            return True
        except Exception as e:
            print(f"转换失败: {str(e)}")
            return False

    def extract_text_from_pdf(self, input_pdf):
        """从PDF中提取文本的示例方法"""
        # 这里是提取PDF文本的逻辑
        # 返回提取的文本
        return "示例文本"  # 替换为实际提取的文本

    def batch_convert(self, pdf_files, output_dir):
        """批量转换PDF文件"""
        results = []
        total_files = len(pdf_files)
        
        for i, pdf_file in enumerate(pdf_files):
            try:
                filename = os.path.basename(pdf_file)
                output_file = os.path.join(output_dir, 
                    os.path.splitext(filename)[0] + '.docx')
                
                success = self.convert_pdf_to_docx(pdf_file, output_file)
                results.append((pdf_file, success))
                
                self.update_progress((i + 1) / total_files * 100)
                
            except Exception as e:
                results.append((pdf_file, False))
                
        return results

    def ocr_with_multiple_engines(self, image):
        """使用多个OCR引擎进行识别"""
        text_results = []
        
        # 预处理图像
        processed_image = self.preprocess_image(image)

        # 1. TrOCR（高精度模型）
        if self.trocr_available:
            try:
                trocr_text = self.ocr_with_trocr(processed_image)
                if trocr_text.strip():
                    text_results.append({
                        'text': trocr_text,
                        'engine': 'TrOCR',
                        'confidence': 0.9  # TrOCR通常具有较高的准确率
                    })
            except Exception as e:
                print(f"TrOCR错误: {str(e)}")

        # 2. EasyOCR with GPU
        if self.easyocr_available:
            try:
                results = self.reader.readtext(np.array(processed_image))
                if results:
                    easy_text = '\n'.join([result[1] for result in results])
                    if easy_text.strip():
                        text_results.append({
                            'text': easy_text,
                            'engine': 'EasyOCR',
                            'confidence': sum(result[2] for result in results) / len(results)
                        })
            except Exception as e:
                print(f"EasyOCR错误: {str(e)}")

        # 3. Tesseract OCR
        tesseract_text = pytesseract.image_to_string(processed_image, lang='chi_sim+eng')
        if tesseract_text.strip():
            text_results.append({
                'text': tesseract_text,
                'engine': 'Tesseract',
                'confidence': 0.5
            })

        # 选择最佳结果
        if text_results:
            best_result = max(text_results, key=lambda x: x['confidence'])
            return best_result['text']

        return ""

    def run_inference_with_onnx(self, image):
        """使用ONNX Runtime运行推理"""
        try:
            # 图像预处理
            if isinstance(image, Image.Image):
                image = np.array(image)
            
            # 转换为模型需要的格式
            # 注意：这里的具体预处理步骤需要根据你的模型要求来调整
            processed_image = cv2.resize(image, (224, 224))  # 示例尺寸
            processed_image = processed_image.astype(np.float32) / 255.0
            processed_image = np.transpose(processed_image, (2, 0, 1))
            processed_image = np.expand_dims(processed_image, axis=0)
            
            # 运行推理
            ort_inputs = {self.ort_session.get_inputs()[0].name: processed_image}
            ort_outputs = self.ort_session.run(None, ort_inputs)
            
            # 处理输出
            # 注意：这里的后处理步骤需要根据你的模型输出格式来调整
            result = ort_outputs[0]
            
            return result
            
        except Exception as e:
            print(f"ONNX推理错误: {str(e)}")
            return None

class PDFConverterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF转文档工具")
        self.root.geometry("600x400")
        self.converter = PDFConverter()
        self.converter.set_progress_callback(self.update_progress)
        self.setup_ui()
    
    def update_progress(self, value):
        """更新进度条"""
        if hasattr(self, 'progress'):
            self.progress['value'] = value
            self.root.update_idletasks()
    
    def setup_ui(self):
        """设置GUI界面"""
        # 创建主框架
        self.main_frame = ttk.Frame(self.root, padding="10")
        self.main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 输入文件选择
        self.input_frame = ttk.LabelFrame(self.main_frame, text="选择PDF文件", padding="5")
        self.input_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        self.input_path = tk.StringVar()
        self.input_entry = ttk.Entry(self.input_frame, textvariable=self.input_path, width=50)
        self.input_entry.grid(row=0, column=0, padx=5)
        
        self.input_button = ttk.Button(self.input_frame, text="浏览", command=self.select_input_file)
        self.input_button.grid(row=0, column=1, padx=5)
        
        # 输出文件选择
        self.output_frame = ttk.LabelFrame(self.main_frame, text="保存位置", padding="5")
        self.output_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        
        self.output_path = tk.StringVar()
        self.output_entry = ttk.Entry(self.output_frame, textvariable=self.output_path, width=50)
        self.output_entry.grid(row=0, column=0, padx=5)
        
        self.output_button = ttk.Button(self.output_frame, text="浏览", command=self.select_output_file)
        self.output_button.grid(row=0, column=1, padx=5)
        
        # 转换按钮
        self.convert_button = ttk.Button(self.main_frame, text="开始转换", command=self.convert_file)
        self.convert_button.grid(row=2, column=0, columnspan=2, pady=20)
        
        # 进度条
        self.progress = ttk.Progressbar(self.main_frame, length=400, mode='determinate')
        self.progress.grid(row=3, column=0, columnspan=2, pady=5)
        
        # 状态标签
        self.status_var = tk.StringVar(value="准备就绪")
        self.status_label = ttk.Label(self.main_frame, textvariable=self.status_var)
        self.status_label.grid(row=4, column=0, columnspan=2, pady=5)

    def select_input_file(self):
        """选择输入文件"""
        filename = filedialog.askopenfilename(
            title="选择PDF文件",
            filetypes=[("PDF文件", "*.pdf")]
        )
        if filename:
            self.input_path.set(filename)
            # 自动设置输出文件名
            output_filename = os.path.splitext(filename)[0] + ".docx"
            self.output_path.set(output_filename)

    def select_output_file(self):
        """选择输出文件"""
        filename = filedialog.asksaveasfilename(
            title="保存文件",
            filetypes=[("Word文档", "*.docx")],
            defaultextension=".docx"
        )
        if filename:
            self.output_path.set(filename)

    def convert_file(self):
        """转换文件"""
        input_file = self.input_path.get()
        output_file = self.output_path.get()
        
        if not input_file or not output_file:
            messagebox.showerror("错误", "请选择输入和输出文件！")
            return
        
        self.progress['value'] = 0
        self.status_var.set("正在转换...")
        self.convert_button['state'] = 'disabled'
        self.root.update()
        
        try:
            success = self.converter.convert_pdf_to_docx(input_file, output_file)
            
            if success:
                self.status_var.set("转换完成！")
                messagebox.showinfo("成功", "文件转换成功！")
            else:
                self.status_var.set("转换失败！")
                messagebox.showerror("错误", "文件转换失败！")
        except Exception as e:
            self.status_var.set("转换错误！")
            messagebox.showerror("错误", f"转换过程中出现错误：{str(e)}")
        finally:
            self.convert_button['state'] = 'normal'

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFConverterGUI(root)
    root.mainloop() 